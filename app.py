import io
import re
import streamlit as st
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from pypdf import PdfReader
from google import genai

# Configuración de página
st.set_page_config(
    page_title="Generador de Sesiones MINEDU - Nivel Inicial",
    page_icon="🎒",
    layout="wide"
)

st.title("🎒 Generador Automático de Sesiones de Aprendizaje - MINEDU")
st.markdown("Sube el Proyecto de Aprendizaje y descarga las sesiones en formato Word (.docx) por cada día.")

# Barra lateral para API Key
with st.sidebar:
    st.header("⚙️ Configuración")
    api_key = st.text_input("Ingresa tu Gemini API Key:", type="password")
    st.markdown("[Obtener API Key gratis en Google AI Studio](https://aistudio.google.com/)")

# Funciones de extracción de texto
def extract_text_from_file(uploaded_file):
    if uploaded_file.name.endswith('.pdf'):
        reader = PdfReader(uploaded_file)
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        return text
    elif uploaded_file.name.endswith('.docx'):
        doc = Document(uploaded_file)
        return "\n".join([p.text for p in doc.paragraphs])
    elif uploaded_file.name.endswith('.txt'):
        return uploaded_file.read().decode('utf-8')
    return ""

# Función para convertir Markdown/Tablas a archivo Word (.docx)
def markdown_to_docx(md_text, title="Sesion_de_Aprendizaje"):
    doc = Document()
    
    # Márgenes de la página
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    lines = md_text.split('\n')
    in_table = False
    table_data = []

    for line in lines:
        stripped = line.strip()

        # Detección de tablas Markdown
        if stripped.startswith('|') and stripped.endswith('|'):
            in_table = True
            # Omitir líneas separadoras de markdown tipo |---|---|
            if re.match(r'^\|[\s\:\-|\+]+\|$', stripped):
                continue
            cells = [c.strip() for c in stripped.split('|')[1:-1]]
            table_data.append(cells)
            continue
        else:
            if in_table and table_data:
                # Renderizar la tabla acumulada
                cols_count = max(len(row) for row in table_data)
                table = doc.add_table(rows=len(table_data), cols=cols_count)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.style = 'Table Grid'

                for row_idx, row in enumerate(table_data):
                    for col_idx, cell_value in enumerate(row):
                        if col_idx < cols_count:
                            cell = table.cell(row_idx, col_idx)
                            # Limpiar formato básico markdown
                            clean_val = cell_value.replace('**', '').replace('<br>', '\n').replace('<br/>', '\n')
                            cell.text = clean_val
                            
                            # Formato para el encabezado de la tabla
                            if row_idx == 0:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.font.bold = True
                                        run.font.size = Pt(10)
                            else:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.font.size = Pt(9.5)

                doc.add_paragraph() # Espaciado
                in_table = False
                table_data = []

        # Títulos
        if stripped.startswith('# '):
            p = doc.add_heading(stripped[2:], level=1)
            p.runs[0].font.color.rgb = RGBColor(0, 51, 102)
        elif stripped.startswith('## '):
            p = doc.add_heading(stripped[3:], level=2)
            p.runs[0].font.color.rgb = RGBColor(0, 102, 153)
        elif stripped.startswith('### '):
            p = doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith('* ') or stripped.startswith('- '):
            clean_item = stripped[2:].replace('**', '')
            doc.add_paragraph(clean_item, style='List Bullet')
        elif stripped:
            clean_p = stripped.replace('**', '')
            doc.add_paragraph(clean_p)

    # Procesar tabla si quedó al final del texto
    if in_table and table_data:
        cols_count = max(len(row) for row in table_data)
        table = doc.add_table(rows=len(table_data), cols=cols_count)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'
        for row_idx, row in enumerate(table_data):
            for col_idx, cell_value in enumerate(row):
                if col_idx < cols_count:
                    cell = table.cell(row_idx, col_idx)
                    cell.text = cell_value.replace('**', '').replace('<br>', '\n')

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Carga de archivos
col1, col2 = st.columns(2)
with col1:
    proyecto_file = st.file_uploader("1. Sube el Proyecto de Aprendizaje (PDF, DOCX o TXT)", type=['pdf', 'docx', 'txt'])
with col2:
    plantilla_file = st.file_uploader("2. Sube la Plantilla / Modelo de Sesión (Opcional)", type=['pdf', 'docx', 'txt'])

if st.button("🚀 Analizar Proyecto y Generar Sesiones", type="primary"):
    if not api_key:
        st.error("Por favor, ingresa tu Gemini API Key en la barra lateral.")
    elif not proyecto_file:
        st.error("Por favor, sube el archivo de tu Proyecto de Aprendizaje.")
    else:
        try:
            client = genai.Client(api_key=api_key)
            
            with st.spinner("Leyendo archivos y analizando la estructura del proyecto..."):
                proyecto_txt = extract_text_from_file(proyecto_file)
                plantilla_txt = extract_text_from_file(plantilla_file) if plantilla_file else "Usar estructura estándar MINEDU para nivel inicial (4 y 5 años)."

                # Paso 1: Consultar número de días
                prompt_dias = f"""
                Analiza el siguiente Proyecto de Aprendizaje de Educación Inicial y responde ÚNICAMENTE con un número entero que represente la cantidad total de días/sesiones programadas en el proyecto.
                
                PROYECTO:
                {proyecto_txt[:4000]}
                """
                response_dias = client.models.generate_content(
                    model="gemini-2.5-flash",
                    contents=prompt_dias
                )
                
                try:
                    num_dias = int(re.search(r'\d+', response_dias.text).group())
                except:
                    num_dias = 5 # Valor por defecto si no detecta número exacto

            st.success(f"Se detectaron **{num_dias} días/sesiones** en el Proyecto de Aprendizaje. Iniciando generación...")

            # Generar cada día e ir mostrando botones de descarga
            for dia in range(1, num_dias + 1):
                with st.spinner(f"Generando Sesión del Día {dia} de {num_dias}..."):
                    prompt_sesion = f"""
                    Actúa como una Especialista y Docente Experta en Educación Inicial (4 y 5 años) del MINEDU (Perú).
                    
                    OBJETIVO:
                    Redacta la Sesión de Aprendizaje COMPLETA para el DÍA {dia} del Proyecto de Aprendizaje adjunto, respetando la secuencia del proyecto y el área correspondiente.
                    
                    MODELO / PLANTILLA DE REFERENCIA:
                    {plantilla_txt}
                    
                    PROYECTO DE APRENDIZAJE ADJUNTO:
                    {proyecto_txt}
                    
                    REGLAS OBLIGATORIAS:
                    1. Información completa: Sin resúmenes, sin 'etc.', sin 'completar aquí'. Incluye diálogos literales de la docente, respuestas esperadas de los niños y secuencia detallada.
                    2. Aplica los procesos didácticos correspondientes al área del día según CNEB-MINEDU.
                    3. Utiliza tablas de Markdown formateadas limpiamente para la estructura de la sesión (Datos Informativos, Propósitos y Evidencias, Preparación, Secuencia Didáctica e Instrumento de Evaluación).
                    """
                    
                    res_sesion = client.models.generate_content(
                        model="gemini-2.5-flash",
                        contents=prompt_sesion
                    )
                    
                    session_md = res_sesion.text
                    docx_buffer = markdown_to_docx(session_md, title=f"Sesion_Dia_{dia}")
                    
                    # Mostrar card con botón de descarga para cada día
                    with st.expander(f"📌 Sesión Día {dia} - Lista para descargar", expanded=True):
                        st.download_button(
                            label=f"📥 Descargar Sesión Día {dia} (.docx)",
                            data=docx_buffer,
                            file_name=f"Sesion_Dia_{dia}_Inicial.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key=f"btn_dia_{dia}"
                        )
            
            st.balloons()
            st.success("¡Todas las sesiones han sido generadas correctamente!")

        except Exception as e:
            st.error(f"Ocurrió un error durante la generación: {str(e)}")