# utils/docx_generator.py
import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from io import BytesIO

def set_cell_background(cell, fill_hex):
    """Aplica color de fondo a una celda de la tabla."""
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Establece márgenes internos (padding) a las celdas."""
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table):
    """Aplica bordes finos y grises a toda la tabla."""
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'<w:top w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/>'
            f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/>'
            f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="E0E0E0"/>'
            f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="E0E0E0"/>'
            f'<w:left w:val="none"/>'
            f'<w:right w:val="none"/>'
            f'</w:tblBorders>'
        )
        tblPr[0].append(borders)

def create_docx_from_text(content: str, title: str) -> BytesIO:
    doc = docx.Document()

    # Configuración de márgenes de página
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Estilo predeterminado de fuente
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Título Principal
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run(title.upper())
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(16)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)  # Azul Oscuro MINEDU

    lines = content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].strip()

        # Omitir líneas vacías o separadores simples
        if not line or line == '---':
            i += 1
            continue

        # Encabezados
        if line.startswith('#'):
            level = min(line.count('#'), 3)
            header_text = line.lstrip('#').strip()
            h = doc.add_heading(header_text, level=level)
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after = Pt(4)
            for run in h.runs:
                run.font.name = 'Arial'
                run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
            i += 1
            continue

        # Detección y creación de Tablas Markdown (| ... |)
        if line.startswith('|') and line.endswith('|'):
            table_data = []
            while i < len(lines) and lines[i].strip().startswith('|') and lines[i].strip().endswith('|'):
                current_line = lines[i].strip()
                # Ignorar la línea divisoria del Markdown (| :--- | :--- |)
                if not set(current_line.replace('|', '').replace(':', '').replace('-', '').strip()).issubset({''}):
                    row_cells = [cell.strip() for cell in current_line.split('|')[1:-1]]
                    table_data.append(row_cells)
                i += 1

            if table_data:
                num_rows = len(table_data)
                num_cols = max(len(row) for row in table_data)

                table = doc.add_table(rows=num_rows, cols=num_cols)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                set_table_borders(table)

                for r_idx, row_content in enumerate(table_data):
                    for c_idx, cell_value in enumerate(row_content):
                        if c_idx < num_cols:
                            cell = table.cell(r_idx, c_idx)
                            cell.text = cell_value.replace('<br>', '\n')
                            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)

                            # Formato especial para el encabezado (Primera fila)
                            if r_idx == 0:
                                set_cell_background(cell, "F2F4F7")  # Gris Claro / Azulado
                                for p in cell.paragraphs:
                                    for r in p.runs:
                                        r.font.bold = True
                                        r.font.size = Pt(9.5)
                                        r.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
                            else:
                                for p in cell.paragraphs:
                                    for r in p.runs:
                                        r.font.size = Pt(9)
            continue

        # Párrafos convencionales
        p = doc.add_paragraph(line.replace('**', ''))
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        i += 1

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer